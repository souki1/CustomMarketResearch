"""
Report export engine.

- blocks -> DOCX  (docxtpl template + python-docx fine edits)
- blocks -> PDF   (reportlab)
"""

from __future__ import annotations

import io
from datetime import datetime, timezone
from pathlib import Path
from typing import Any
from xml.sax.saxutils import escape

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from docxtpl import DocxTemplate
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_RIGHT
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import (
    HRFlowable,
    Image as RLImage,
    ListFlowable,
    ListItem,
    Paragraph,
    Preformatted,
    SimpleDocTemplate,
    Spacer,
    Table,
    TableStyle,
)

_DIR = Path(__file__).resolve().parent
TEMPLATE_PATH = _DIR / "templates" / "report_template.docx"

_ALIGN_MAP = {
    "left": WD_ALIGN_PARAGRAPH.LEFT,
    "center": WD_ALIGN_PARAGRAPH.CENTER,
    "right": WD_ALIGN_PARAGRAPH.RIGHT,
}

_CALLOUT_COLORS: dict[str, RGBColor] = {
    "amber": RGBColor(0xB4, 0x5B, 0x09),
    "blue": RGBColor(0x1D, 0x4E, 0xD8),
    "emerald": RGBColor(0x04, 0x7A, 0x57),
    "slate": RGBColor(0x47, 0x55, 0x69),
}

_CALLOUT_BG: dict[str, str] = {
    "amber": "FFF7ED",
    "blue": "EFF6FF",
    "emerald": "ECFDF5",
    "slate": "F8FAFC",
}


def _ensure_template() -> Path:
    if TEMPLATE_PATH.exists():
        return TEMPLATE_PATH
    from reports.create_template import build_template
    return build_template()


def _align(block: dict[str, Any]) -> int:
    return _ALIGN_MAP.get(block.get("align", "left"), WD_ALIGN_PARAGRAPH.LEFT)


def _set_paragraph_shading(paragraph, hex_color: str) -> None:
    """Apply background shading to a paragraph via XML."""
    shading_elm = paragraph._element.makeelement(
        qn("w:shd"),
        {
            qn("w:val"): "clear",
            qn("w:color"): "auto",
            qn("w:fill"): hex_color,
        },
    )
    pPr = paragraph._element.get_or_add_pPr()
    pPr.append(shading_elm)


def _add_block(doc: Document, block: dict[str, Any]) -> None:
    """Append one report block to the document using python-docx."""
    btype = block.get("type", "paragraph")
    align = _align(block)

    if btype == "title":
        p = doc.add_paragraph()
        p.alignment = align
        run = p.add_run(block.get("text", ""))
        run.font.size = Pt(24)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

    elif btype == "heading":
        h = doc.add_heading(block.get("text", ""), level=1)
        h.alignment = align

    elif btype == "subheading":
        h = doc.add_heading(block.get("text", ""), level=2)
        h.alignment = align

    elif btype == "paragraph":
        p = doc.add_paragraph(block.get("text", ""))
        p.alignment = align

    elif btype == "bullets":
        for item in block.get("items", []):
            p = doc.add_paragraph(item, style="List Bullet")
            p.alignment = align

    elif btype == "numbered":
        for item in block.get("items", []):
            p = doc.add_paragraph(item, style="List Number")
            p.alignment = align

    elif btype == "divider":
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        style = block.get("style", "solid")
        run = p.add_run("— " * 20 if style == "dashed" else "─" * 50)
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)

    elif btype == "callout":
        tone = block.get("tone", "amber")
        p = doc.add_paragraph()
        p.alignment = align
        p.paragraph_format.left_indent = Inches(0.3)
        p.paragraph_format.right_indent = Inches(0.3)
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(6)
        _set_paragraph_shading(p, _CALLOUT_BG.get(tone, "FFF7ED"))
        run = p.add_run(block.get("text", ""))
        run.font.color.rgb = _CALLOUT_COLORS.get(tone, _CALLOUT_COLORS["amber"])
        run.font.size = Pt(10)

    elif btype == "quote":
        p = doc.add_paragraph()
        p.alignment = align
        p.paragraph_format.left_indent = Inches(0.5)
        p.style = doc.styles["Quote"] if "Quote" in [s.name for s in doc.styles] else doc.styles["Normal"]
        run = p.add_run(block.get("text", ""))
        run.font.italic = True
        run.font.color.rgb = RGBColor(0x64, 0x64, 0x64)

    elif btype == "spacer":
        doc.add_paragraph("")

    elif btype == "image":
        src = block.get("src", "")
        caption = block.get("caption", "") or block.get("alt", "")
        if src and Path(src).exists():
            doc.add_picture(src, width=Inches(4.5))
        if caption:
            p = doc.add_paragraph(caption)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.runs[0].font.size = Pt(9)
            p.runs[0].font.italic = True
            p.runs[0].font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    elif btype == "metric":
        p = doc.add_paragraph()
        p.alignment = align
        label_run = p.add_run(f"{block.get('label', 'Metric')}: ")
        label_run.font.size = Pt(11)
        label_run.font.bold = True
        label_run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
        value_run = p.add_run(block.get("value", "—"))
        value_run.font.size = Pt(14)
        value_run.font.bold = True
        value_run.font.color.rgb = RGBColor(0x6D, 0x28, 0xD9)

    elif btype == "code":
        p = doc.add_paragraph()
        p.alignment = align
        p.paragraph_format.left_indent = Inches(0.2)
        _set_paragraph_shading(p, "F4F4F5")
        run = p.add_run(block.get("text", ""))
        run.font.name = "Consolas"
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

    elif btype == "table":
        rows_data: list[list[Any]] = block.get("rows", [])
        if not rows_data:
            return
        show_header = block.get("showHeader", True)
        n_cols = max(len(r) for r in rows_data) if rows_data else 1
        n_rows = len(rows_data)

        table = doc.add_table(rows=n_rows, cols=n_cols)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.style = "Table Grid"

        for r_idx, row_data in enumerate(rows_data):
            row_cells = table.rows[r_idx].cells
            for c_idx in range(n_cols):
                cell = row_cells[c_idx]
                raw = row_data[c_idx] if c_idx < len(row_data) else ""
                if isinstance(raw, dict) and raw.get("type") == "link":
                    label = str(raw.get("label", "Link")).strip()
                    cell.text = label
                    if cell.paragraphs:
                        for p in cell.paragraphs:
                            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                else:
                    cell.text = "" if raw is None else str(raw)

        if show_header and n_rows > 0:
            for cell in table.rows[0].cells:
                for p in cell.paragraphs:
                    for run in p.runs:
                        run.font.bold = True
                        run.font.size = Pt(10)

    else:
        p = doc.add_paragraph(str(block.get("text", "")))
        p.alignment = align


def render_docx_bytes(
    title: str,
    blocks: list[dict[str, Any]],
) -> bytes:
    """
    Render report blocks into a .docx file in memory.

    1. Load the docxtpl template and render title / date.
    2. Remove the body marker paragraph.
    3. Append each block with python-docx fine edits.
    """
    tpl_path = _ensure_template()
    tpl = DocxTemplate(str(tpl_path))

    now = datetime.now(timezone.utc).strftime("%B %d, %Y")
    tpl.render({"title": title, "generated_date": now, "__body_marker__": ""})

    doc: Document = tpl.docx  # underlying python-docx Document

    # Remove the marker paragraph left by the template
    for para in doc.paragraphs:
        if "__body_marker__" in para.text:
            p_element = para._element
            p_element.getparent().remove(p_element)
            break

    for block in blocks:
        _add_block(doc, block)

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def _pdf_alignment(block: dict[str, Any]) -> int:
    align = str(block.get("align", "left")).lower()
    if align == "center":
        return TA_CENTER
    if align == "right":
        return TA_RIGHT
    return TA_LEFT


def _pdf_para_style(base: ParagraphStyle, name: str, align: int) -> ParagraphStyle:
    return ParagraphStyle(name=name, parent=base, alignment=align)


def _pdf_text(value: Any) -> str:
    # Escape HTML because reportlab Paragraph supports a subset of markup.
    text = "" if value is None else str(value)
    return escape(text).replace("\n", "<br/>")


def _add_pdf_block(
    story: list[Any],
    block: dict[str, Any],
    styles: dict[str, ParagraphStyle],
) -> None:
    btype = str(block.get("type", "paragraph"))
    align = _pdf_alignment(block)

    if btype == "title":
        story.append(Paragraph(_pdf_text(block.get("text", "")), _pdf_para_style(styles["title"], "title_align", align)))
        story.append(Spacer(1, 0.16 * inch))
        return

    if btype == "heading":
        story.append(Paragraph(_pdf_text(block.get("text", "")), _pdf_para_style(styles["h1"], "h1_align", align)))
        story.append(Spacer(1, 0.10 * inch))
        return

    if btype == "subheading":
        story.append(Paragraph(_pdf_text(block.get("text", "")), _pdf_para_style(styles["h2"], "h2_align", align)))
        story.append(Spacer(1, 0.08 * inch))
        return

    if btype == "paragraph":
        story.append(Paragraph(_pdf_text(block.get("text", "")), _pdf_para_style(styles["body"], "body_align", align)))
        story.append(Spacer(1, 0.08 * inch))
        return

    if btype in {"bullets", "numbered"}:
        items = [str(i) for i in block.get("items", []) if str(i).strip()]
        if items:
            bullet_type = "bullet" if btype == "bullets" else "1"
            list_items = [
                ListItem(Paragraph(_pdf_text(item), _pdf_para_style(styles["body"], "list_item", align)), leftIndent=8)
                for item in items
            ]
            story.append(ListFlowable(list_items, bulletType=bullet_type, start="1"))
            story.append(Spacer(1, 0.08 * inch))
        return

    if btype == "divider":
        story.append(HRFlowable(width="100%", thickness=0.8, color=colors.HexColor("#D4D4D8")))
        story.append(Spacer(1, 0.10 * inch))
        return

    if btype == "callout":
        tone = str(block.get("tone", "amber"))
        fg = _CALLOUT_COLORS.get(tone, _CALLOUT_COLORS["amber"])
        bg_hex = _CALLOUT_BG.get(tone, "FFF7ED")
        color_hex = f"#{fg[0]:02X}{fg[1]:02X}{fg[2]:02X}"
        text = _pdf_text(block.get("text", ""))
        para = Paragraph(
            f"<para backColor='#{bg_hex}' textColor='{color_hex}'>{text}</para>",
            _pdf_para_style(styles["body"], "callout_align", align),
        )
        story.append(para)
        story.append(Spacer(1, 0.10 * inch))
        return

    if btype == "quote":
        story.append(Paragraph(_pdf_text(block.get("text", "")), _pdf_para_style(styles["quote"], "quote_align", align)))
        story.append(Spacer(1, 0.10 * inch))
        return

    if btype == "spacer":
        story.append(Spacer(1, 0.18 * inch))
        return

    if btype == "image":
        src = str(block.get("src", "")).strip()
        caption = str(block.get("caption", "") or block.get("alt", "")).strip()
        if src:
            src_path = Path(src)
            if src_path.exists():
                try:
                    image = RLImage(str(src_path))
                    image._restrictSize(6.0 * inch, 4.5 * inch)
                    story.append(image)
                    story.append(Spacer(1, 0.06 * inch))
                except Exception:
                    # Ignore broken image files and continue export.
                    pass
        if caption:
            story.append(Paragraph(_pdf_text(caption), _pdf_para_style(styles["caption"], "caption_align", TA_CENTER)))
            story.append(Spacer(1, 0.08 * inch))
        return

    if btype == "metric":
        label = _pdf_text(block.get("label", "Metric"))
        value = _pdf_text(block.get("value", "—"))
        story.append(Paragraph(f"<b>{label}:</b> <font color='#6D28D9'><b>{value}</b></font>", _pdf_para_style(styles["body"], "metric_align", align)))
        story.append(Spacer(1, 0.08 * inch))
        return

    if btype == "code":
        story.append(Preformatted(str(block.get("text", "")), styles["code"]))
        story.append(Spacer(1, 0.10 * inch))
        return

    if btype == "table":
        rows_data: list[list[Any]] = block.get("rows", [])
        if rows_data:
            n_cols = max((len(r) for r in rows_data), default=1)
            cell_style = ParagraphStyle(
                name="TableCellInner",
                parent=styles["body"],
                fontSize=9,
                leading=11,
                textColor=colors.HexColor("#374151"),
            )
            normalized: list[list[Any]] = []
            show_hdr = bool(block.get("showHeader", True))
            for r_idx, row in enumerate(rows_data):
                next_row: list[Any] = []
                is_header_row = show_hdr and r_idx == 0
                for c_idx in range(n_cols):
                    cell = row[c_idx] if c_idx < len(row) else ""
                    if isinstance(cell, dict) and cell.get("type") == "link":
                        href = str(cell.get("href", "")).strip()
                        label = str(cell.get("label", "Link")).strip() or "Link"
                        if href:
                            href_esc = (
                                href.replace("&", "&amp;")
                                .replace('"', "&quot;")
                                .replace("<", "&lt;")
                                .replace(">", "&gt;")
                            )
                            label_esc = _pdf_text(label)
                            inner = f'<a href="{href_esc}" color="#2563EB">{label_esc}</a>'
                            if is_header_row:
                                inner = f"<b>{inner}</b>"
                            next_row.append(Paragraph(inner, cell_style))
                        else:
                            t = _pdf_text(label)
                            next_row.append(Paragraph(f"<b>{t}</b>" if is_header_row else t, cell_style))
                    else:
                        t = _pdf_text(cell)
                        next_row.append(Paragraph(f"<b>{t}</b>" if is_header_row else t, cell_style))
                normalized.append(next_row)
            col_widths = block.get("colWidths")
            kw: dict[str, Any] = {"hAlign": "CENTER"}
            if isinstance(col_widths, list) and len(col_widths) == n_cols:
                try:
                    kw["colWidths"] = [float(w) * inch for w in col_widths]
                except (TypeError, ValueError):
                    pass
            tbl = Table(normalized, **kw)
            style = TableStyle([
                ("GRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#D4D4D8")),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("FONTNAME", (0, 0), (-1, -1), "Helvetica"),
                ("FONTSIZE", (0, 0), (-1, -1), 9),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
                ("TOPPADDING", (0, 0), (-1, -1), 5),
            ])
            if block.get("showHeader", True) and normalized:
                style.add("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#F4F4F5"))
                style.add("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold")
            tbl.setStyle(style)
            story.append(tbl)
            story.append(Spacer(1, 0.12 * inch))
        return

    # Fallback to plain paragraph for unknown block types.
    story.append(Paragraph(_pdf_text(block.get("text", "")), _pdf_para_style(styles["body"], "fallback_align", align)))
    story.append(Spacer(1, 0.08 * inch))


def render_pdf_bytes(
    title: str,
    blocks: list[dict[str, Any]],
) -> bytes:
    """
    Generate PDF directly from blocks using reportlab (in memory).
    """
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(
        buffer,
        pagesize=A4,
        leftMargin=0.75 * inch,
        rightMargin=0.75 * inch,
        topMargin=0.75 * inch,
        bottomMargin=0.75 * inch,
        title=title or "Report",
    )

    sample = getSampleStyleSheet()
    styles: dict[str, ParagraphStyle] = {
        "title": ParagraphStyle(
            name="Title",
            parent=sample["Title"],
            fontName="Helvetica-Bold",
            fontSize=22,
            leading=26,
            textColor=colors.HexColor("#1A1A2E"),
            spaceAfter=8,
        ),
        "h1": ParagraphStyle(
            name="Heading1",
            parent=sample["Heading1"],
            fontName="Helvetica-Bold",
            fontSize=16,
            leading=20,
            textColor=colors.HexColor("#111827"),
            spaceBefore=6,
            spaceAfter=4,
        ),
        "h2": ParagraphStyle(
            name="Heading2",
            parent=sample["Heading2"],
            fontName="Helvetica-Bold",
            fontSize=13,
            leading=16,
            textColor=colors.HexColor("#1F2937"),
            spaceBefore=4,
            spaceAfter=3,
        ),
        "body": ParagraphStyle(
            name="Body",
            parent=sample["BodyText"],
            fontName="Helvetica",
            fontSize=10.5,
            leading=14,
            textColor=colors.HexColor("#374151"),
        ),
        "caption": ParagraphStyle(
            name="Caption",
            parent=sample["BodyText"],
            fontName="Helvetica-Oblique",
            fontSize=9,
            leading=12,
            textColor=colors.HexColor("#6B7280"),
            alignment=TA_CENTER,
        ),
        "quote": ParagraphStyle(
            name="Quote",
            parent=sample["BodyText"],
            fontName="Helvetica-Oblique",
            fontSize=10.5,
            leading=14,
            leftIndent=14,
            textColor=colors.HexColor("#52525B"),
        ),
        "code": ParagraphStyle(
            name="Code",
            parent=sample["Code"],
            fontName="Courier",
            fontSize=9,
            leading=11,
            backColor=colors.HexColor("#F4F4F5"),
            borderPadding=6,
            borderWidth=0.5,
            borderColor=colors.HexColor("#E5E7EB"),
            borderRadius=2,
        ),
    }

    story: list[Any] = []
    # If blocks already include a title block, this will still render consistently.
    now = datetime.now(timezone.utc).strftime("%B %d, %Y")
    story.append(Paragraph(_pdf_text(title or "Untitled report"), styles["title"]))
    story.append(Paragraph(_pdf_text(f"Generated on {now}"), ParagraphStyle(
        name="GeneratedDate",
        parent=styles["body"],
        fontSize=9,
        textColor=colors.HexColor("#6B7280"),
    )))
    story.append(Spacer(1, 0.18 * inch))

    for block in blocks:
        _add_pdf_block(story, block, styles)

    doc.build(story)
    return buffer.getvalue()
