"""
Generate the docxtpl base template (report_template.docx).

Run once:  python -m reports.create_template
The template uses Jinja2 tags rendered by docxtpl at export time.
"""

from pathlib import Path

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

TEMPLATE_DIR = Path(__file__).resolve().parent / "templates"
TEMPLATE_PATH = TEMPLATE_DIR / "report_template.docx"


def build_template() -> Path:
    doc = Document()

    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(11)
    font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    style.paragraph_format.space_after = Pt(4)

    for level in range(1, 4):
        heading_style = doc.styles[f"Heading {level}"]
        heading_style.font.name = "Calibri"
        heading_style.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1.2)
    section.right_margin = Inches(1.2)

    # Title placeholder rendered by docxtpl
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = title_para.add_run("{{ title }}")
    run.font.size = Pt(26)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

    # Date line
    date_para = doc.add_paragraph()
    date_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = date_para.add_run("Generated: {{ generated_date }}")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
    run.font.italic = True

    # Thin rule
    doc.add_paragraph("_______________________________________________")

    # Body placeholder — export engine will replace this paragraph
    marker = doc.add_paragraph()
    marker.add_run("{{r __body_marker__ }}")

    TEMPLATE_DIR.mkdir(parents=True, exist_ok=True)
    doc.save(str(TEMPLATE_PATH))
    return TEMPLATE_PATH


if __name__ == "__main__":
    path = build_template()
    print(f"Template written to {path}")
